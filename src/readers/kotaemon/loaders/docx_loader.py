import sys
import unicodedata
from pathlib import Path
from typing import List, Optional

# Add the parent directory to the system path for imports
sys.path.append(str(Path(__file__).parent.parent.parent.parent))

import pandas as pd
from llama_index.core.readers.base import BaseReader
from src.readers.kotaemon.base import Document, split_text


class DocxReader(BaseReader):
    def __init__(self, max_words_per_page: int = 2048, *args, **kwargs):
        self._import_docx()
        self.max_words_per_page = max_words_per_page

    def _import_docx(self):
        try:
            import docx  # noqa
        except ImportError:
            raise ImportError(
                "docx is not installed. "
                "Please install it using `pip install python-docx`"
            )

    def _load_single_table(self, table) -> List[List[str]]:
        return [[cell.text for cell in row.cells] for row in table.rows]

    def load_data(
        self, file_path: Path, extra_info: Optional[dict] = None, **kwargs
    ) -> List[Document]:
        import docx  # Import here to minimize global imports

        file_path = Path(file_path).resolve()
        doc = docx.Document(str(file_path))

        all_text = "\n".join(
            [unicodedata.normalize("NFKC", p.text) for p in doc.paragraphs]
        )
        pages = split_text(all_text, max_tokens=self.max_words_per_page)

        tables = [pd.DataFrame({a[0]: a[1:] for a in self._load_single_table(t)}) for t in doc.tables]

        extra_info = extra_info or {}
        documents = [
            Document(
                text=table.to_csv(index=False).strip(),
                metadata={
                    "table_origin": table.to_csv(index=False),
                    "type": "table",
                    **extra_info,
                },
                metadata_template="",
                metadata_separator="",
            )
            for table in tables
        ]

        documents.extend(
            [
                Document(
                    text=non_table_text.strip(),
                    metadata={"page_label": i + 1, **extra_info},
                )
                for i, non_table_text in enumerate(pages)
            ]
        )

        return documents
